"""
Resume parser - Extract text from PDF and DOCX files.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Optional
from loguru import logger


@dataclass
class ResumeText:
    """Extracted resume text and sections."""
    full_text: str
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    summary: Optional[str] = None
    skills: list[str] = None
    experience: list[str] = None
    education: list[str] = None
    projects: list[str] = None

    def __post_init__(self):
        if self.skills is None:
            self.skills = []
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []
        if self.projects is None:
            self.projects = []


class ResumeParser:
    """Parse resume from PDF or DOCX files."""

    @staticmethod
    def parse(file_path: str, file_type: Optional[str] = None) -> ResumeText:
        """
        Parse resume from file and extract text and sections.

        Args:
            file_path: Path to resume file
            file_type: 'pdf' or 'docx' (auto-detected if None)

        Returns:
            ResumeText object with extracted content

        Raises:
            ValueError: If file type unsupported or parsing fails
        """
        path = Path(file_path)

        # Auto-detect file type if not provided
        if file_type is None:
            file_type = path.suffix.lower().lstrip('.')

        if file_type == 'pdf':
            return ResumeParser._parse_pdf(file_path)
        elif file_type == 'docx':
            return ResumeParser._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}. Use 'pdf' or 'docx'.")

    @staticmethod
    def _parse_pdf(file_path: str) -> ResumeText:
        """
        Parse PDF resume using pdfplumber.

        Args:
            file_path: Path to PDF file

        Returns:
            ResumeText object
        """
        try:
            import pdfplumber
        except ImportError:
            raise ValueError("pdfplumber not installed. Install with: pip install pdfplumber")

        try:
            full_text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        full_text += text + "\n"

            logger.info(f"Extracted {len(full_text)} characters from PDF")

            # Parse sections
            resume = ResumeText(full_text=full_text)
            ResumeParser._extract_sections(full_text, resume)

            return resume

        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def _parse_docx(file_path: str) -> ResumeText:
        """
        Parse DOCX resume using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            ResumeText object
        """
        try:
            from docx import Document
        except ImportError:
            raise ValueError("python-docx not installed. Install with: pip install python-docx")

        try:
            doc = Document(file_path)
            full_text = ""

            # Extract all paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text += row_text + "\n"

            logger.info(f"Extracted {len(full_text)} characters from DOCX")

            # Parse sections
            resume = ResumeText(full_text=full_text)
            ResumeParser._extract_sections(full_text, resume)

            return resume

        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    @staticmethod
    def _extract_sections(text: str, resume: ResumeText) -> None:
        """
        Extract structured sections from resume text.

        Args:
            text: Full resume text
            resume: ResumeText object to populate
        """
        lines = text.split('\n')

        # Try to find contact info (usually first few lines)
        contact_section = ' '.join(lines[:10])

        # Extract email
        import re
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', contact_section)
        if email_match:
            resume.email = email_match.group()

        # Extract phone (common patterns)
        phone_match = re.search(r'(\+?\d{1,3}[-.\s]?)?\(?(\d{3})\)?[-.\s]?(\d{3})[-.\s]?(\d{4})', contact_section)
        if phone_match:
            resume.phone = phone_match.group(0)

        # Extract name (usually first line, longest name candidate)
        for line in lines[:5]:
            line = line.strip()
            if line and len(line) < 100 and len(line.split()) >= 2:
                # Likely a name if it has 2+ words and is short
                resume.name = line
                break

        # Find skill mentions (commonly formatted with keywords)
        text_lower = text.lower()
        skill_keywords = [
            'python', 'javascript', 'typescript', 'java', 'c++', 'c#', 'golang', 'rust', 'ruby', 'php',
            'react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'express', 'spring',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'devops',
            'sql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
            'machine learning', 'tensorflow', 'pytorch', 'nlp', 'computer vision',
            'git', 'agile', 'scrum', 'rest', 'graphql', 'microservices',
            'html', 'css', 'sass', 'tailwind', 'figma', 'ui/ux'
        ]

        found_skills = []
        for skill in skill_keywords:
            if skill in text_lower and skill not in found_skills:
                found_skills.append(skill)

        resume.skills = found_skills

        # Extract experience mentions (line with "years" or job titles)
        experience_keywords = ['experience', 'role', 'manager', 'developer', 'engineer', 'analyst', 'coordinator']
        for line in lines:
            if any(kw in line.lower() for kw in experience_keywords) and len(line.strip()) > 10:
                if line.strip() not in resume.experience and len(resume.experience) < 10:
                    resume.experience.append(line.strip())

        # Extract education mentions
        education_keywords = ['bachelor', 'master', 'phd', 'b.s.', 'm.s.', 'university', 'college', 'degree']
        for line in lines:
            if any(kw in line.lower() for kw in education_keywords) and len(line.strip()) > 5:
                if line.strip() not in resume.education and len(resume.education) < 5:
                    resume.education.append(line.strip())

        logger.info(f"Extracted {len(resume.skills)} skills, {len(resume.experience)} experience items, {len(resume.education)} education items")
